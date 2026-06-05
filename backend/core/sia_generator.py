"""
SIA (Security Impact Analysis) document generator.
Produces a draft SIA as both a structured dict (for UI display)
and a downloadable DOCX.
"""

import json
import os
from datetime import datetime, timezone
from pathlib import Path
from dataclasses import dataclass

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.mapping_engine import MappingResult


@dataclass
class SIADocument:
    system_name: str
    change_reference: str
    prepared_date: str
    tier: str
    tier_color: str
    change_description: str
    system_impact_summary: str
    affected_controls_analysis: list[dict]
    affected_artifacts: list[dict]
    risk_determination: str
    risk_level: str
    recommended_actions: list[str]
    ao_action_required: str
    raw_json: dict


def _build_sia_prompt(mapping: MappingResult, system_name: str, change_ref: str) -> str:
    changes_text = "\n".join(
        f"- {c.change_type_name} (confidence: {c.confidence}, source: {c.source})"
        + (f"\n  Evidence: {', '.join(c.evidence[:2])}" if c.evidence else "")
        for c in mapping.detected_changes
    )

    artifacts_text = "\n".join(
        f"- {a.id}: {a.name} (urgency: {a.urgency})"
        for a in mapping.affected_artifacts
    )

    controls_text = "\n".join(
        f"- {c.id}: {c.title} ({c.family} family)"
        for c in mapping.affected_controls
    )

    return f"""You are a federal cybersecurity compliance specialist writing a Security Impact Analysis (SIA) document.

System: {system_name}
Change Reference: {change_ref}
FedRAMP Tier: {mapping.tier['name']}
AO Action Required: {mapping.tier['ao_action']}

Detected Changes:
{changes_text}

Affected Artifacts:
{artifacts_text}

Affected NIST 800-53 Controls:
{controls_text}

Write a draft SIA with professional federal compliance language. Be specific — reference actual change types, control IDs, and artifacts by name. Do not use boilerplate.

Respond with JSON only — no markdown fences:
{{
  "change_description": "2-3 sentences describing what changed and why it matters for security",
  "system_impact_summary": "2-3 sentences on the overall security posture impact",
  "affected_controls_analysis": [
    {{
      "control_id": "SC-7",
      "control_title": "Boundary Protection",
      "impact_description": "specific impact on this control",
      "current_implementation": "what the SSP currently says (general)",
      "required_update": "what needs to change in the SSP"
    }}
  ],
  "risk_determination": "2-3 sentences on residual risk after recommended actions",
  "risk_level": "Low|Moderate|High",
  "additional_recommended_actions": ["specific actionable item 1", "specific actionable item 2"]
}}"""


def generate_sia_content(
    mapping: MappingResult,
    system_name: str = "[SYSTEM NAME]",
    change_ref: str = "N/A",
    api_key: str | None = None,
) -> SIADocument:
    """Generate SIA content using Claude."""
    key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
    now = datetime.now(timezone.utc).strftime("%B %d, %Y")

    if key:
        client = anthropic.Anthropic(api_key=key)
        prompt = _build_sia_prompt(mapping, system_name, change_ref)

        message = client.messages.create(
            model="claude-opus-4-7",
            max_tokens=2048,
            system="You are a FedRAMP compliance expert. Respond with valid JSON only.",
            messages=[{"role": "user", "content": prompt}],
        )

        try:
            data = json.loads(message.content[0].text)
        except json.JSONDecodeError:
            data = _fallback_sia_content(mapping)
    else:
        data = _fallback_sia_content(mapping)

    # Merge recommended actions
    all_actions = mapping.recommended_actions + data.get("additional_recommended_actions", [])

    return SIADocument(
        system_name=system_name,
        change_reference=change_ref,
        prepared_date=now,
        tier=mapping.tier["name"],
        tier_color=mapping.tier["color"],
        change_description=data.get("change_description", ""),
        system_impact_summary=data.get("system_impact_summary", ""),
        affected_controls_analysis=data.get("affected_controls_analysis", []),
        affected_artifacts=[
            {"id": a.id, "name": a.name, "urgency": a.urgency, "phase": a.phase}
            for a in mapping.affected_artifacts
        ],
        risk_determination=data.get("risk_determination", ""),
        risk_level=data.get("risk_level", "Moderate"),
        recommended_actions=list(dict.fromkeys(all_actions)),  # deduplicate, preserve order
        ao_action_required=mapping.tier["ao_action"],
        raw_json=data,
    )


def _fallback_sia_content(mapping: MappingResult) -> dict:
    """Fallback content when Claude API is unavailable."""
    change_names = [c.change_type_name for c in mapping.detected_changes]
    controls = [c.id for c in mapping.affected_controls]
    return {
        "change_description": f"The system underwent the following changes: {', '.join(change_names)}. These changes affect the security posture and require assessment per agency policy.",
        "system_impact_summary": f"This change is classified as {mapping.tier['name']} and affects {mapping.total_control_count} NIST 800-53 controls across {mapping.total_artifact_count} authorization artifacts.",
        "affected_controls_analysis": [
            {"control_id": c.id, "control_title": c.title, "impact_description": f"This control is affected by the detected {c.triggered_by[0]} change.", "current_implementation": "[Review current SSP narrative]", "required_update": "[Update narrative to reflect change]"}
            for c in mapping.affected_controls[:5]
        ],
        "risk_determination": "Risk determination pending ISSO review. Recommended actions should be completed before this SIA is finalized.",
        "risk_level": "Moderate",
        "additional_recommended_actions": [],
    }


def build_sia_docx(sia: SIADocument) -> bytes:
    """Build and return a DOCX SIA document as bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("SECURITY IMPACT ANALYSIS (SIA)")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("FOR OFFICIAL USE ONLY — PRE-DECISIONAL DRAFT").font.size = Pt(9)

    doc.add_paragraph()

    # System identification table
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    fields = [
        ("System Name", sia.system_name),
        ("FISMA ID / System ID", "[PLACEHOLDER]"),
        ("Change Reference", sia.change_reference),
        ("Date Prepared", sia.prepared_date),
        ("Prepared By", "cATO Advisor (AI-assisted draft — requires ISSO review)"),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Tier banner
    tier_para = doc.add_paragraph()
    tier_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tier_run = tier_para.add_run(f"CHANGE TIER: {sia.tier.upper()}  |  AO ACTION: {sia.ao_action_required.upper()}")
    tier_run.bold = True
    tier_run.font.size = Pt(12)
    color_map = {"green": (0,128,0), "yellow": (180,130,0), "orange": (200,80,0), "red": (180,0,0)}
    r, g, b = color_map.get(sia.tier_color, (0,0,0))
    tier_run.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph()

    def add_section(num: str, title: str):
        p = doc.add_paragraph()
        run = p.add_run(f"{num}. {title}")
        run.bold = True
        run.font.size = Pt(12)

    # Section 1
    add_section("1", "System Identification")
    doc.add_paragraph("See header table above.")

    # Section 2
    add_section("2", "Change Description")
    doc.add_paragraph(sia.change_description or "[AI-generated description pending]")

    doc.add_paragraph("Detected Change Types:")
    for change in sia.affected_controls_analysis:
        pass  # handled below

    # Section 3
    add_section("3", "Security Impact Assessment")
    doc.add_paragraph(sia.system_impact_summary or "[AI-generated summary pending]")

    # Section 4
    add_section("4", "Affected Controls Analysis")
    for ctrl in sia.affected_controls_analysis:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{ctrl.get('control_id', '')} — {ctrl.get('control_title', '')}").bold = True
        doc.add_paragraph(f"Impact: {ctrl.get('impact_description', '')}")
        doc.add_paragraph(f"Required SSP Update: {ctrl.get('required_update', '')}")

    # Section 5
    add_section("5", "Affected Artifacts")
    art_table = doc.add_table(rows=1 + len(sia.affected_artifacts), cols=4)
    art_table.style = "Table Grid"
    headers = ["Artifact ID", "Name", "Phase", "Urgency"]
    for i, h in enumerate(headers):
        art_table.rows[0].cells[i].text = h
        art_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, art in enumerate(sia.affected_artifacts, 1):
        row = art_table.rows[i]
        row.cells[0].text = art["id"]
        row.cells[1].text = art["name"]
        row.cells[2].text = art["phase"]
        row.cells[3].text = art["urgency"].upper()

    doc.add_paragraph()

    # Section 6
    add_section("6", "Risk Determination")
    risk_p = doc.add_paragraph()
    risk_p.add_run(f"Risk Level: {sia.risk_level}").bold = True
    doc.add_paragraph(sia.risk_determination or "[Pending ISSO review]")

    # Section 7
    add_section("7", "Recommended Actions")
    for action in sia.recommended_actions:
        doc.add_paragraph(action, style="List Bullet")

    # Section 8 — Signatures
    add_section("8", "Approval Signatures")
    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.style = "Table Grid"
    sig_headers = ["Role", "Name / Signature", "Date"]
    for i, h in enumerate(sig_headers):
        sig_table.rows[0].cells[i].text = h
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, role in enumerate(["ISSO", "ISSM", "Authorizing Official (AO)"], 1):
        sig_table.rows[i].cells[0].text = role
        sig_table.rows[i].cells[1].text = ""
        sig_table.rows[i].cells[2].text = ""

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("NOTE: ").bold = True
    note.add_run("This document was AI-assisted. All sections require ISSO review and validation before submission. AI-generated content is a draft starting point, not a final determination.")
    note.runs[-1].font.size = Pt(9)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
